import re
from py_youtube import Data
from docx import Document
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import Formatter, TextFormatter
from deepmultilingualpunctuation import PunctuationModel
from semantic_text_splitter import TextSplitter

# Get the transcript for a video
# 1. Get the video
# 2. Retrieve the video ID
# 3. Get the transcript for the video
# 4. Format the transcript
# 5. Save the transcript to a file

def get_video_link():
    """
    Get the video link from the user.
    Return: video link as a string
    """
    try:
        return input('Enter the video link: ')
    except Exception as e:
        print(f'Error getting video link: {e}')
        return None

def extract_video_data(video_link):
    """
    Extract video data from the video link.
    Args:
        video_link (str): The link of the video.
    Return: 
        video_id (str): The ID of the video.
        video_title (str): The title of the video.
    """
    try:
        video_data = Data(video_link).data()
        video_id = video_data['id']
        video_title = video_data['title']
        return video_id, video_title
    except Exception as e:
        print(f'Error extracting video data: {e}')
        return None, None

def get_video():
    """
    Get the video link from the user and extract the video data.
    Return: 
        video_id (str): The ID of the video.
        video_title (str): The title of the video.
    """
    video_link = get_video_link()
    return extract_video_data(video_link)

def format_transcript(transcript):
    """
    Format the transcript.
    Args:
        transcript (str): The transcript to format.
    Return: 
        formatted_transcript (str): The formatted transcript.
    """
    model = PunctuationModel()
    formatter = TextFormatter()
    text_format = formatter.format_transcript(transcript)
    text_format = model.restore_punctuation(text_format)
    text_format = re.subn(r"(^|[.!?])\s*([a-z])", lambda p: p.group(0).upper(), text_format)[0]

    # Replace instances of '[Music]' with an empty string
    text_format = re.sub(r'\[Music\]', '', text_format)

    return text_format

def split_transcript(formatted_transcript):
    """
    Split the formatted transcript into chunks.
    Args:
        formatted_transcript (str): The formatted transcript to split.
    Return: 
        transcript_chunks (list): The transcript split into chunks.
    """
    max_tokens = 250
    splitter = TextSplitter.from_tiktoken_model("gpt-3.5-turbo", max_tokens)
    transcript_chunks = splitter.chunks(formatted_transcript)
    return transcript_chunks

def get_transcript(video_id):
    """
    Get the transcript for the video.
    Args:
        video_id (str): The ID of the video.
    Return: 
        transcript_chunks (list): The transcript for the video, split into chunks.
    """
    try:
        # Get the transcript for the video
        transcript = YouTubeTranscriptApi.get_transcript(video_id)

        # Format the transcript
        formatted_transcript = format_transcript(transcript)

        # Split the transcript into chunks
        transcript_chunks = split_transcript(formatted_transcript)

        return transcript_chunks
    except Exception as e:
        print(f'Error getting transcript: {e}')
        return None

def save_to_file(video_title, transcript):
    """
    Save the transcript to a file
    Return: Docx File
    """
    try:
        video_title = re.sub(r'[\\/:*?"<>|]', '', video_title)
        filename = f'{video_title}.docx'
        doc = Document()
        doc.add_heading(f"Transcript for '{video_title}'", level=1)
        for paragraph in transcript:
            doc.add_paragraph(paragraph)
            doc.add_paragraph()
        doc.save(filename)
        print(f"Transcript saved to '{filename}'")
    except Exception as e:
        print(f'Error saving transcript to file: {e}')



if __name__ == '__main__':
    # Get the video
    video_id, video_title = get_video()
    # Get the transcript for the video
    transcript = get_transcript(video_id)
    # Save the transcript to a file
    save_to_file(video_title, transcript)
#video = 'https://youtu.be/d-yXXMarULk?si=wk4Um0WimqpB4IKk'
